"""Text extraction service. Handles PDF, DOCX, and TXT files.
Includes vision OCR fallback for scanned/image-based PDFs using Groq API."""

import base64
from pathlib import Path

import fitz  # PyMuPDF
import httpx
from docx import Document as DocxDocument

from config import settings
from core.exceptions import ExtractionError
from core.logging import get_logger
from dataclasses import dataclass, field

logger = get_logger(__name__)


@dataclass
class PageContent:
    """Content of a single page."""
    page_num: int
    text: str
    char_offset: int


@dataclass
class ExtractedDocument:
    """Result of text extraction."""
    text: str
    pages: list[PageContent]
    page_count: int
    char_count: int
    metadata: dict = field(default_factory=dict)


async def extract_with_vision(file_path: Path) -> ExtractedDocument:
    """Extract text from scanned/image-based PDFs using Groq vision model.

    Converts each PDF page to an image and uses the vision model for OCR.
    """
    if not settings.groq_api_key:
        raise ExtractionError(
            "This PDF appears to be scanned. Vision OCR requires GROQ_API_KEY to be configured."
        )

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ExtractionError(f"Failed to open PDF for vision OCR: {e}")

    if doc.needs_pass:
        doc.close()
        raise ExtractionError("This PDF is password-protected and cannot be processed.")

    all_text = []
    pages = []
    char_offset = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        # Convert page to image at 2x zoom for quality
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = pix.tobytes("png")
        base64_img = base64.b64encode(img_data).decode("utf-8")
        data_url = f"data:image/png;base64,{base64_img}"

        # Call Groq vision for OCR
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    settings.groq_api_url,
                    headers={
                        "Authorization": f"Bearer {settings.groq_api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": settings.groq_model,
                        "messages": [{
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "Extract all text from this document page. "
                                        "Preserve formatting, tables, and structure. "
                                        "Output only the extracted text with no commentary."
                                    ),
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {"url": data_url},
                                },
                            ],
                        }],
                        "temperature": 0.1,
                        "max_completion_tokens": 4096,
                    },
                )
                response.raise_for_status()
                data = response.json()
                page_text = data["choices"][0]["message"]["content"]
        except httpx.TimeoutException:
            logger.warning("vision_ocr_timeout", page=page_num + 1)
            page_text = f"[OCR TIMEOUT - Page {page_num + 1}]"
        except Exception as e:
            logger.error("vision_ocr_failed", page=page_num + 1, error=str(e))
            page_text = f"[OCR FAILED - Page {page_num + 1}]"

        all_text.append(page_text)
        pages.append(PageContent(
            page_num=page_num + 1,
            text=page_text,
            char_offset=char_offset,
        ))
        char_offset += len(page_text)

    doc.close()

    full_text = "\n\n".join(all_text)

    logger.info(
        "vision_ocr_complete",
        page_count=len(pages),
        char_count=len(full_text),
    )

    return ExtractedDocument(
        text=full_text,
        pages=pages,
        page_count=len(pages),
        char_count=len(full_text),
        metadata={"extraction_method": "vision_ocr"},
    )


def extract_pdf(file_path: Path) -> ExtractedDocument:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ExtractionError(f"Failed to open PDF: {e}")

    if doc.needs_pass:
        doc.close()
        raise ExtractionError("This PDF is password-protected and cannot be processed.")

    pages = []
    full_text_parts = []
    char_offset = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()

        # Check if page has extractable text
        if not text.strip():
            # Try to get text from blocks
            blocks = page.get_text("blocks")
            if blocks:
                text = " ".join(b[4] for b in blocks if b[4].strip())

        pages.append(PageContent(
            page_num=page_num + 1,
            text=text,
            char_offset=char_offset,
        ))
        full_text_parts.append(text)
        char_offset += len(text)

    doc.close()

    full_text = "\n".join(full_text_parts)

    if not full_text.strip():
        raise ExtractionError(
            "This PDF appears to be scanned. Text extraction requires OCR which is not enabled."
        )

    return ExtractedDocument(
        text=full_text,
        pages=pages,
        page_count=len(pages),
        char_count=len(full_text),
        metadata={"page_count": len(pages)},
    )


def extract_docx(file_path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise ExtractionError(f"Failed to open DOCX: {e}")

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        raise ExtractionError("The DOCX file appears to be empty or contains no extractable text.")

    # DOCX doesn't have pages, so we treat section breaks as page-like divisions
    pages = [PageContent(page_num=0, text=full_text, char_offset=0)]

    return ExtractedDocument(
        text=full_text,
        pages=pages,
        page_count=1,
        char_count=len(full_text),
        metadata={"paragraph_count": len(paragraphs)},
    )


def extract_txt(file_path: Path) -> ExtractedDocument:
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception as e:
        raise ExtractionError(f"Failed to read TXT file: {e}")

    if not text.strip():
        raise ExtractionError("The TXT file appears to be empty or contains only whitespace.")

    pages = [PageContent(page_num=0, text=text, char_offset=0)]

    return ExtractedDocument(
        text=text,
        pages=pages,
        page_count=1,
        char_count=len(text),
        metadata={},
    )


async def extract(file_path: Path, file_type: str) -> ExtractedDocument:
    """Route to the correct extractor based on file type.

    For PDFs, tries standard text extraction first. If the PDF appears
    to be scanned (no extractable text), falls back to vision OCR using
    the Groq vision model.
    """
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "txt": extract_txt,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ExtractionError(f"Unsupported file type: {file_type}")

    if file_type.lower() == "pdf":
        # Try standard text extraction first
        try:
            result = extract_pdf(file_path)
            # Check if we actually got any text
            if result.text.strip():
                return result
        except ExtractionError:
            pass

        # Fallback to vision OCR
        logger.info("pdf_text_extraction_failed_falling_back_to_vision_ocr", path=str(file_path))
        return await extract_with_vision(file_path)

    return extractor(file_path)
